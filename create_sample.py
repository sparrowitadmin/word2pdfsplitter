#!/usr/bin/env python3
"""
Script to create a sample Word document for testing the PDF splitter.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_document():
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Document for PDF Splitter', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'This is a sample document created for testing the Word to PDF Splitter application. '
        'It contains multiple sections across several pages to demonstrate the splitting functionality.'
    )
    doc.add_paragraph(
        'You can use this document to test various page ranges and configurations. '
        'Each section below represents different content that might be split into separate PDF files.'
    )
    
    # Add page break
    doc.add_page_break()
    
    # Section 1
    doc.add_heading('Section 1: Getting Started', level=1)
    doc.add_paragraph(
        'This is the first section of the document. It provides basic information about '
        'how to use the application and what to expect.'
    )
    doc.add_paragraph(
        'Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor '
        'incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud '
        'exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.'
    )
    
    doc.add_heading('Key Features:', level=2)
    features = [
        'Easy file upload with drag and drop support',
        'Automatic Word to PDF conversion',
        'Custom page range selection',
        'Multiple output files from single document',
        'User-friendly web interface'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Add page break
    doc.add_page_break()
    
    # Section 2
    doc.add_heading('Section 2: Technical Details', level=1)
    doc.add_paragraph(
        'This section covers the technical aspects of the application, including the '
        'technologies used and system requirements.'
    )
    
    doc.add_heading('Technologies Used:', level=2)
    doc.add_paragraph('Backend: Python, Flask')
    doc.add_paragraph('PDF Processing: pypdf')
    doc.add_paragraph('Word Processing: python-docx, docx2pdf')
    doc.add_paragraph('Frontend: HTML5, CSS3, JavaScript')
    
    doc.add_paragraph(
        '\nDuis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore '
        'eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in '
        'culpa qui officia deserunt mollit anim id est laborum.'
    )
    
    # Add page break
    doc.add_page_break()
    
    # Section 3
    doc.add_heading('Section 3: Usage Examples', level=1)
    doc.add_paragraph(
        'Here are some practical examples of how to use the Word to PDF Splitter application.'
    )
    
    doc.add_heading('Example 1: Simple Split', level=2)
    doc.add_paragraph(
        'Upload a 10-page document and split it into two files: '
        'pages 1-5 as "part1.pdf" and pages 6-10 as "part2.pdf".'
    )
    
    doc.add_heading('Example 2: Multiple Sections', level=2)
    doc.add_paragraph(
        'Take a 20-page report and create separate PDFs for the introduction (pages 1-3), '
        'main content (pages 4-17), and conclusion (pages 18-20).'
    )
    
    doc.add_heading('Example 3: Chapter Extraction', level=2)
    doc.add_paragraph(
        'Extract individual chapters from a book manuscript, creating a separate PDF '
        'for each chapter based on page ranges.'
    )
    
    # Add page break
    doc.add_page_break()
    
    # Section 4
    doc.add_heading('Section 4: Best Practices', level=1)
    doc.add_paragraph(
        'Follow these best practices to get the most out of the PDF splitter application.'
    )
    
    best_practices = [
        'Always verify the total page count before configuring splits',
        'Use descriptive file names for better organization',
        'Create the output folder before starting the process',
        'Test with a small document first',
        'Check that page ranges don\'t overlap if you want distinct sections',
        'Keep the original Word file as a backup'
    ]
    
    for practice in best_practices:
        doc.add_paragraph(practice, style='List Number')
    
    doc.add_paragraph(
        '\nSed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium '
        'doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore '
        'veritatis et quasi architecto beatae vitae dicta sunt explicabo.'
    )
    
    # Add page break
    doc.add_page_break()
    
    # Section 5
    doc.add_heading('Section 5: Troubleshooting', level=1)
    doc.add_paragraph(
        'Common issues and their solutions.'
    )
    
    doc.add_heading('Issue: Conversion Failed', level=2)
    doc.add_paragraph(
        'Make sure you have Microsoft Word (Windows/Mac) or LibreOffice (Linux) installed '
        'on your system. The docx2pdf library requires one of these applications to perform '
        'the conversion.'
    )
    
    doc.add_heading('Issue: Invalid Page Range', level=2)
    doc.add_paragraph(
        'Double-check that your start page is less than or equal to your end page, and '
        'that both values are within the total page count of the document.'
    )
    
    doc.add_heading('Issue: Output Folder Not Found', level=2)
    doc.add_paragraph(
        'Ensure the output folder exists before starting the process. The application does '
        'not automatically create folders that don\'t exist.'
    )
    
    # Add page break
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'Thank you for using the Word to PDF Splitter application! This tool is designed '
        'to make document management easier by allowing you to extract specific pages from '
        'your Word documents as separate PDF files.'
    )
    doc.add_paragraph(
        'We hope you find this application useful for your document processing needs. '
        'Feel free to experiment with different configurations and page ranges to achieve '
        'your desired results.'
    )
    
    doc.add_paragraph(
        '\nNemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, '
        'sed quia consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt. '
        'Neque porro quisquam est, qui dolorem ipsum quia dolor sit amet, consectetur, '
        'adipisci velit.'
    )
    
    doc.add_paragraph()
    final = doc.add_paragraph('Happy splitting! 📄✂️📑')
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    doc.save('sample_document.docx')
    print("✅ Sample document created successfully: sample_document.docx")
    print(f"📄 The document contains approximately {len(doc.paragraphs)} paragraphs across 6+ pages")

if __name__ == '__main__':
    create_sample_document()

